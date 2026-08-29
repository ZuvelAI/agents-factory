from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from agents_factory.modules.knowledge.ingestion.contracts import (
    ExtractedBlock,
    ExtractedDocument,
    FetchedSource,
    IngestionRejected,
)


class DocxExtractor:
    def extract(self, fetched: FetchedSource) -> ExtractedDocument:
        try:
            document = Document(BytesIO(fetched.content))
        except (PackageNotFoundError, ValueError):
            raise IngestionRejected("docx_malformed") from None
        blocks: list[ExtractedBlock] = []
        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if text:
                blocks.append(
                    ExtractedBlock(
                        kind="TEXT",
                        text=text,
                        locator={
                            **fetched.locator,
                            "paragraph": index,
                            "style": (
                                None
                                if paragraph.style is None
                                else paragraph.style.name
                            ),
                        },
                    )
                )
        for index, table in enumerate(document.tables, start=1):
            rows = tuple(
                tuple(cell.text.strip() for cell in row.cells) for row in table.rows
            )
            text = "\n".join(" | ".join(row) for row in rows).strip()
            if text:
                blocks.append(
                    ExtractedBlock(
                        kind="TABLE",
                        text=text,
                        rows=rows,
                        locator={**fetched.locator, "table": index},
                    )
                )
        if not blocks:
            raise IngestionRejected("source_has_no_extractable_text")
        return ExtractedDocument(
            title=fetched.filename,
            blocks=tuple(blocks),
            source_digest=fetched.content_digest,
        )
