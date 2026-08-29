from __future__ import annotations

import hashlib
from io import BytesIO
from pathlib import Path
from uuid import UUID

import httpx
import pytest
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from agents_factory.modules.knowledge.ingestion.contracts import (
    FetchedSource,
    IngestionRejected,
    SourceDescriptor,
)
from agents_factory.modules.knowledge.ingestion.docx import DocxExtractor
from agents_factory.modules.knowledge.ingestion.normalizer import KnowledgeNormalizer
from agents_factory.modules.knowledge.ingestion.pdf import PdfExtractor
from agents_factory.modules.knowledge.ingestion.spreadsheet import SpreadsheetExtractor
from agents_factory.modules.knowledge.ingestion.website import (
    WebsiteExtractor,
    WebsiteFetcher,
)
from agents_factory.modules.knowledge.models import KnowledgeSourceType


FIXTURES = Path(__file__).resolve().parents[3] / "fixtures/knowledge"
TENANT_ID = UUID("10000000-0000-0000-0000-000000000018")
SOURCE_ID = UUID("20000000-0000-0000-0000-000000000018")


def source(
    source_type: KnowledgeSourceType,
    configuration: dict[str, object],
) -> SourceDescriptor:
    return SourceDescriptor(
        tenant_id=TENANT_ID,
        source_id=SOURCE_ID,
        source_type=source_type,
        authority="AUTHORITATIVE",
        configuration=configuration,
    )


def fetched(
    *, source_type: str, content: bytes, media_type: str, filename: str
) -> FetchedSource:
    return FetchedSource(
        descriptor=source(source_type, {}),
        content=content,
        media_type=media_type,
        filename=filename,
        locator={"fixture": filename},
        content_digest=hashlib.sha256(content).hexdigest(),
    )


@pytest.mark.asyncio
async def test_website_fetch_is_allowlisted_public_and_sanitized() -> None:
    html = (FIXTURES / "site/index.html").read_bytes()
    client = httpx.AsyncClient(
        transport=httpx.MockTransport(
            lambda request: httpx.Response(
                200,
                content=html,
                headers={"content-type": "text/html; charset=utf-8"},
                request=request,
            )
        )
    )

    async def public_resolver(_hostname: str) -> tuple[str, ...]:
        return ("93.184.216.34",)

    try:
        result = await WebsiteFetcher(
            allowed_hosts=frozenset({"example.com"}),
            client=client,
            resolver=public_resolver,
        ).fetch(source("WEBSITE", {"url": "https://example.com/"}))
    finally:
        await client.aclose()
    extracted = WebsiteExtractor().extract(result)

    assert "Atención con cita previa." in extracted.blocks[0].text
    assert "secret()" not in extracted.blocks[0].text

    async def private_resolver(_hostname: str) -> tuple[str, ...]:
        return ("169.254.169.254",)

    with pytest.raises(IngestionRejected, match="private_network"):
        await WebsiteFetcher(
            allowed_hosts=frozenset({"metadata.example"}),
            resolver=private_resolver,
        ).fetch(source("WEBSITE", {"url": "https://metadata.example/"}))


def test_docx_spreadsheet_and_encrypted_pdf_are_bounded_and_typed() -> None:
    document_stream = BytesIO()
    document = Document()
    document.add_heading("Manual", level=1)
    document.add_paragraph("Procedimiento aprobado")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Servicio"
    table.cell(0, 1).text = "Precio"
    table.cell(1, 0).text = "Consulta"
    table.cell(1, 1).text = "100"
    document.save(document_stream)
    docx = DocxExtractor().extract(
        fetched(
            source_type="DOCX",
            content=document_stream.getvalue(),
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            filename="manual.docx",
        )
    )
    assert {block.kind for block in docx.blocks} == {"TEXT", "TABLE"}

    workbook_stream = BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Catálogo"
    sheet.append(("servicio", "precio"))
    sheet.append(("consulta", 100))
    workbook.save(workbook_stream)
    spreadsheet = SpreadsheetExtractor().extract(
        fetched(
            source_type="SPREADSHEET",
            content=workbook_stream.getvalue(),
            media_type=(
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ),
            filename="catalog.xlsx",
        )
    )
    assert spreadsheet.blocks[0].locator["sheet"] == "Catálogo"

    pdf_stream = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.encrypt("password")
    writer.write(pdf_stream)
    with pytest.raises(IngestionRejected, match="pdf_encrypted"):
        PdfExtractor().extract(
            fetched(
                source_type="PDF",
                content=pdf_stream.getvalue(),
                media_type="application/pdf",
                filename="policy.pdf",
            )
        )


def test_google_sheet_rows_become_structured_draft_not_vector_only() -> None:
    content = (FIXTURES / "google_sheet_rows.json").read_bytes()
    descriptor = source(
        "SPREADSHEET",
        {
            "fact_kind": "BUSINESS_HOURS",
            "fact_key": "operations.business_hours.main",
        },
    )
    extracted = SpreadsheetExtractor().extract(
        FetchedSource(
            descriptor=descriptor,
            content=content,
            media_type="application/json",
            filename="google_sheet_rows.json",
            locator={"drive_file_id": "sheet-1"},
            content_digest=hashlib.sha256(content).hexdigest(),
        )
    )

    normalized = KnowledgeNormalizer().normalize(
        source=descriptor,
        document=extracted,
    )

    assert len(normalized.facts) == 1
    assert normalized.facts[0].kind == "BUSINESS_HOURS"
    assert normalized.documents == ()
